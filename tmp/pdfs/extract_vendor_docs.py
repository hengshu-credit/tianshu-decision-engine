from pathlib import Path

from docx import Document
from pypdf import PdfReader


PDFS = {
    "bingjian_qingyun24": Path(r"E:\workspace\data-dictionary-collection\data-dictionary-collection-main\三方数据字典\冰鉴\冰鉴个人评分平台青云分24API文档.pdf"),
    "pudao_index_140": Path(r"C:\Users\Administrator\Downloads\朴道海纳履约指数-1.4.0.pdf"),
    "pudao_interface_20230912": Path(r"E:\workspace\data-dictionary-collection\data-dictionary-collection-main\三方数据字典\海纳数科\朴道征信有限公司信息服务接口规格说明202300912.pdf"),
    "tianchuang_xingyao_pro_v2d": Path(r"E:\workspace\data-dictionary-collection\data-dictionary-collection-main\三方数据字典\天创\天创API－综合信息－信用司南_星耀版_Pro_V2d.pdf"),
    "baihang_special_attention_20": Path(r"E:\workspace\data-dictionary-collection\data-dictionary-collection-main\三方数据字典\百行\百行征信数据服务接口规范_特别关注2.0 - V1.0.pdf"),
}

DOCX = Path(r"E:\workspace\data-dictionary-collection\data-dictionary-collection-main\三方数据字典\融360\200727 占信分综合风险分层分v1.docx")
OUTPUT_DIR = Path(__file__).resolve().parent / "extracted"


def extract_pdf(name: str, path: Path) -> None:
    reader = PdfReader(path)
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        parts.append(f"\n===== PAGE {index} =====\n")
        parts.append((page.extract_text() or "").replace("\x00", ""))
    (OUTPUT_DIR / f"{name}.txt").write_text("".join(parts), encoding="utf-8")


def extract_docx(path: Path) -> None:
    document = Document(path)
    parts = ["===== PARAGRAPHS =====\n"]
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text + "\n")
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n===== TABLE {table_index} =====\n")
        for row in table.rows:
            parts.append("\t".join(cell.text.replace("\n", " | ") for cell in row.cells) + "\n")
    (OUTPUT_DIR / "rong360_zhanxin_score_v1.txt").write_text("".join(parts), encoding="utf-8")


OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
for pdf_name, pdf_path in PDFS.items():
    extract_pdf(pdf_name, pdf_path)
extract_docx(DOCX)
print(OUTPUT_DIR)
